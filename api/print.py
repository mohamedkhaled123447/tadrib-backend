from docx import Document
from functools import reduce

days_arabic = ["الأحد", "الاثنين", "الثلاثاء", "الأربعاء", "الخميس", "الجمعة", "السبت"]


def generate_545_doc(
    data,
    subjects,
    len,
    week,
    month,
    totalTrainingHours,
    totalBoundEducationHours,
    start,
    end,
):
    doc_545 = Document("media/545.docx")
    title = doc_545.paragraphs[0]
    title.text = f"مستخرج الموضوعات والساعات عن الأسبوع  {week}  شهر  {month} في الفترة من {start} الي {end}"
    table = doc_545.tables[0]
    for row in data["the545"]:
        row.reverse()
    new_table_row = table.add_row().cells
    new_table_row[8].text = ":اولا:الموضوعات الرئيسية"
    new_table_row = table.add_row().cells
    new_table_row[8].text = "اللياقة"
    new_table_row[7].text = str(sum(totalTrainingHours))
    totalTrainingHours.reverse()
    totalBoundEducationHours.reverse()
    for i in range(7):
        if totalTrainingHours[i] == 0:
            new_table_row[i].text = "-"
        else:
            new_table_row[i].text = str(totalTrainingHours[i])

    for row_id, row in enumerate(data["the545"]):
        if sum(data["mat"][row_id]) == 0:
            continue
        if row_id == len:
            new_table_row = table.add_row().cells
            new_table_row[8].text = ":ثانيا:الموضوعات العامة"
            new_table_row = table.add_row().cells
            new_table_row[8].text = "تعليم اولي"
            new_table_row[7].text = str(sum(totalBoundEducationHours))
            for i in range(7):
                if totalBoundEducationHours[i] == 0:
                    new_table_row[i].text = "-"
                else:
                    new_table_row[i].text = str(totalBoundEducationHours[i])
        new_table_row = table.add_row().cells
        new_table_row[7].text = str(sum(data["mat"][row_id]))
        new_table_row[8].text = subjects[row_id]
        for cell_id, cell in enumerate(row):
            if cell == "":
                new_table_row[cell_id].text = "-"
            else:
                new_table_row[cell_id].text = cell
    new_table_row = table.add_row().cells
    new_table_row[8].text = "الاجمالي"
    column_sums = [reduce(lambda x, y: x + y, col) for col in zip(*data["mat"])]
    column_sums.reverse()
    new_table_row[7].text = str(
        sum(column_sums) + sum(totalTrainingHours) + sum(totalBoundEducationHours)
    )
    for i in range(7):
        new_table_row[i].text = str(
            column_sums[i] + totalTrainingHours[i] + totalBoundEducationHours[i]
        )
    doc_545.save("media/new_545.docx")
